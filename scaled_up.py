import numpy as np
import networkx as nx
import random
from scipy.spatial import Delaunay
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Code 1 
# Feel free to change the 1200 vlaues to tailor the output
def random_points(min_nodes, max_nodes, scale=1200):
    n = random.randint(min_nodes, max_nodes)
    points = np.random.rand(n, 2) * scale
    return points

def delaunay_triangulation(points):
    tri = Delaunay(points)
    G = nx.Graph()
    
    for simplex in tri.simplices:
        u, v, w = simplex
        u_coord, v_coord, w_coord = points[u], points[v], points[w]
        G.add_edge(u, v, weight=np.linalg.norm(u_coord - v_coord))
        G.add_edge(v, w, weight=np.linalg.norm(v_coord - w_coord))
        G.add_edge(w, u, weight=np.linalg.norm(w_coord - u_coord))
    
    return G, points

# Code 2
def process_data(G, pos, document):
    # floyd warshall for the all pair shortest length
    if not nx.is_connected(G):
        Gcc = sorted(nx.connected_components(G), key=len, reverse=True)
        G = G.subgraph(Gcc[0])
    all_pairs_shortest_paths = nx.algorithms.shortest_paths.floyd_warshall_numpy(G)

    # effective resistance of the graph
    total_effective_resistance = 0
    for i, u in enumerate(G.nodes):
        for j, v in enumerate(G.nodes):
            if i < j:
                shortest_path_length = all_pairs_shortest_paths[i][j]
                effective_resistance = 1 / shortest_path_length
                total_effective_resistance += effective_resistance

    # average shortest path
    total_shortest_paths = sum([sum(row) for row in all_pairs_shortest_paths])
    average_shortest_path = total_shortest_paths / (len(G.nodes) * (len(G.nodes) - 1))

    # average edge length
    total_edge_lengths = sum([G.edges[edge]['weight'] for edge in G.edges])
    average_edge_length = total_edge_lengths / len(G.edges)

    # drawing the graph
    nx.draw_networkx_nodes(G, pos, node_size=15)
    nx.draw_networkx_edges(G, pos)
    graph_filename = f'graph_{len(G.nodes)}.png'
    plt.savefig(graph_filename)
    plt.close()

    print(f"Graph with {len(G.nodes)} nodes")
    print("Total effective resistance:", total_effective_resistance)
    print("Average shortest path:", average_shortest_path)
    print("Average edge length:", average_edge_length)

    # inserting the results into the Word document
    document.add_heading(f'Graph with {len(G.nodes)} nodes', level=1)
    document.add_picture(graph_filename)
    document.add_paragraph(f'Number of nodes: {len(G.nodes)}')
    document.add_paragraph(f'Average shortest path: {average_shortest_path}')
    document.add_paragraph(f'Average edge length: {average_edge_length}')
    document.add_paragraph(f'Total effective resistance: {total_effective_resistance}')

    return document

def run_simulation(num_iterations):
    document = Document()
    document.add_heading('Graph Analysis', level=1)

    for i in range(num_iterations):
        points = random_points(60, 100)
        G, pos = delaunay_triangulation(points)
        document = process_data(G, pos, document)

    document.save(r'C:/Users/adnan/OneDrive/Documents/Complexity_Shokkom_Sharma/test104.docx')

if __name__ == "__main__":
    num_iterations = 5  # Define the number of iterations here
    run_simulation(num_iterations)

       
