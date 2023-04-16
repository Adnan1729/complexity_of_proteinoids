import pandas as pd
import networkx as nx
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def process_data(sheet_name, document):
    # read data from excel sheet
    df = pd.read_excel(filename, sheet_name=sheet_name)

    G = nx.Graph()

    # adding nodes
    for _, row in df.iterrows():
        x1, y1, x2, y2 = row['x1'], row['y1'], row['x2'], row['y2']
        G.add_node((x1, y1))
        G.add_node((x2, y2))

    # adding edges
    for _, row in df.iterrows():
        x1, y1, x2, y2 = row['x1'], row['y1'], row['x2'], row['y2']
        distance = ((x2 - x1)**2 + (y2 - y1)**2)**0.5
        G.add_edge((x1, y1), (x2, y2), weight=distance)

    # floyd warshall for the all pair shortest length
    if not nx.is_connected(G):
        Gcc = sorted(nx.connected_components(G), key=len, reverse=True)
        G = G.subgraph(Gcc[0])
    all_pairs_shortest_paths = nx.algorithms.shortest_paths.floyd_warshall_numpy(G)

    # effective resistance of the graph
    total_effective_resistance = 0
    for i, u in enumerate(G.nodes):
        for j, v in enumerate(G.nodes):
            if i < j:
                shortest_path_length = all_pairs_shortest_paths[i][j]
                effective_resistance = 1 / shortest_path_length
                total_effective_resistance += effective_resistance

    # average shortest path
    total_shortest_paths = sum([sum(row) for row in all_pairs_shortest_paths])
    average_shortest_path = total_shortest_paths / (len(G.nodes) * (len(G.nodes) - 1))

    # average edge length
    total_edge_lengths = sum([G.edges[edge]['weight'] for edge in G.edges])
    average_edge_length = total_edge_lengths / len(G.edges)

    # drawing the graph
    pos = nx.spring_layout(G)
    nx.draw_networkx_nodes(G, pos, node_size=15)
    nx.draw_networkx_edges(G, pos)
    graph_filename = f'graph_{sheet_name}.png'
    plt.savefig(graph_filename)
    plt.close()

    # just checking how things look
    pos = nx.spring_layout(G)
    nx.draw_networkx_nodes(G, pos, node_size=15)
    nx.draw_networkx_edges(G, pos)
    plt.savefig('graph.png')
    plt.show()

    print(f"Sheet: {sheet_name}")
    print("Total effective resistance:", total_effective_resistance)
    print("Average shortest path:", average_shortest_path)
    print("Average edge length:", average_edge_length)

    # inserting the results into the Word document
    document.add_heading(f'Sheet: {sheet_name}', level=1)
    document.add_picture(graph_filename)
    document.add_paragraph(f'Average shortest path: {average_shortest_path}')
    document.add_paragraph(f'Average edge length: {average_edge_length}')
    document.add_paragraph(f'Total effective resistance: {total_effective_resistance}')

    return document


if __name__ == "__main__":
    # set up the input parameters
    filename = '.../Complex
