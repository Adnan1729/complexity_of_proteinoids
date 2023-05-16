import pandas as pd
import networkx as nx
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from scipy.spatial import Delaunay
import random

def delaunay_triangulation(points):
    tri = Delaunay(points)
    G = nx.Graph()
    
    for i, point in enumerate(points):
        G.add_node(i, pos=point)
    
    for simplex in tri.simplices:
        u, v, w = simplex
        u_coord, v_coord, w_coord = points[u], points[v], points[w]
        G.add_edge(u, v, weight=np.linalg.norm(np.subtract(u_coord, v_coord)))
        G.add_edge(v, w, weight=np.linalg.norm(np.subtract(v_coord, w_coord)))
        G.add_edge(w, u, weight=np.linalg.norm(np.subtract(w_coord, u_coord)))
    
    return G, points

def remove_random_edges(G, p=0.9):
    G_copy = G.copy()  # create a copy of the graph
    num_edges = len(G_copy.edges())
    num_to_remove = int(num_edges * p)
    
    edges = list(G_copy.edges())
    random.shuffle(edges)
    
    for edge in edges[:num_to_remove]:
        G_copy.remove_edge(*edge)
        
    return G_copy  # return the modified copy


def percolation_threshold(G, G_removed):
    # calculate the size of the largest connected component before removing edges
    largest_cc_size_before = len(max(nx.connected_components(G), key=len))
    
    # calculate the size of the largest connected component after removing edges
    largest_cc_size_after = len(max(nx.connected_components(G_removed), key=len))
    
    print("Number of edges before removing:", G.number_of_edges())
    print("Number of edges after removing:", G_removed.number_of_edges())
    print("Number of connected components before removing:", nx.number_connected_components(G))
    print("Number of connected components after removing:", nx.number_connected_components(G_removed))
    
    return largest_cc_size_after / largest_cc_size_before

# effective resistance, average shortest path and average edge length

filename = 'C:/Users/adnan/OneDrive/Documents/Complexity_Shokkom_Sharma/img_nodes_2.xlsx'
sheet_names = ['img1','img2', 'img3', 'img4','saksham_img1','saksham_img2'] 

# create a new Word document
document = Document()

##
for sheet_name in sheet_names:
    df = pd.read_excel(filename, sheet_name=sheet_name)

    G = nx.Graph()

    #nodes to graph
    for _, row in df.iterrows():
        x1, y1, x2, y2 = row['x1'], row['y1'], row['x2'], row['y2']
        G.add_node((x1, y1))
        G.add_node((x2, y2))

    #edges to graph
    for _, row in df.iterrows():
        x1, y1, x2, y2 = row['x1'], row['y1'], row['x2'], row['y2']
        distance = ((x2 - x1)**2 + (y2 - y1)**2)**0.5 #euclidean distance between nodes
        G.add_edge((x1, y1), (x2, y2), weight=distance)

    if not nx.is_connected(G):
        Gcc = sorted(nx.connected_components(G), key=len, reverse=True)
        G = G.subgraph(Gcc[0])
    all_pairs_shortest_paths = nx.algorithms.shortest_paths.floyd_warshall_numpy(G)

    # calculate the effective resistance of the graph
    total_effective_resistance = 0
    for i, u in enumerate(G.nodes):
        for j, v in enumerate(G.nodes):
            if i < j:
                shortest_path_length = all_pairs_shortest_paths[i][j]
                effective_resistance = 1 / shortest_path_length
                total_effective_resistance += effective_resistance

    #average shortest path 
    total_shortest_paths = sum([sum(row) for row in all_pairs_shortest_paths])
    average_shortest_path = total_shortest_paths / (len(G.nodes) * (len(G.nodes) - 1))
    
    #average edge length
    total_edge_lengths = sum([G.edges[edge]['weight'] for edge in G.edges])
    average_edge_length = total_edge_lengths / len(G.edges)
    
    #drawing the graph
    pos = nx.spring_layout(G)
    nx.draw_networkx_nodes(G, pos, node_size=15)
    nx.draw_networkx_edges(G, pos)
    graph_filename = f'graph_{sheet_name}.png'
    plt.savefig(graph_filename)
    plt.close()
    
    #-- non-technical parts below. 
    
    #just checking how things look
    pos = nx.spring_layout(G)
    nx.draw_networkx_nodes(G, pos, node_size=15)
    nx.draw_networkx_edges(G, pos)
    plt.savefig('graph.png')
    plt.show()

    print(f"Sheet: {sheet_name}")
    print("Total effective resistance:", total_effective_resistance)
    print("Average shortest path:", average_shortest_path)
    print("Average edge length:", average_edge_length)
    
    # Percolation threshold calculation:
    points = [(x, y) for x, y in G.nodes()]  # Modified to get the nodes from the graph
    G_delaunay, _ = delaunay_triangulation(points)
    G_random_removed = remove_random_edges(G_delaunay, p=0.9)
    threshold = percolation_threshold(G_delaunay, G_random_removed)

    print(f"Percolation threshold: {threshold}")

    # Inserting the results into the Word document
    document.add_paragraph(f'Percolation threshold: {threshold}')
        
    #inserting the results into the Word document
    document.add_heading(f'Sheet: {sheet_name}', level=1)
    document.add_picture(graph_filename)
    document.add_paragraph(f'Average shortest path: {average_shortest_path}')
    document.add_paragraph(f'Average edge length: {average_edge_length}')
    document.add_paragraph(f'Total effective resistance: {total_effective_resistance}')

# save the Word document
document.save('C:/Users/adnan/OneDrive/Documents/Complexity_Shokkom_Sharma/Emperical_Graph_Analysis.docx')


